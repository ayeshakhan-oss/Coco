"""
Extract CV text for Job 41 (Growth Manager - Karachi) new-status applicants.
Batch of 2026-08-10 screening. Saves text to output/cv_texts_job41_new_batch/.

Uses Neon HTTPS SQL API (port 443) — reliable when port 5432 is blocked.
Fetches resume_data one application at a time (payloads up to ~5 MB).
"""

import os, sys, base64, io, json, re
import requests
from dotenv import load_dotenv

import PyPDF2
try:
    import fitz
    import pytesseract
    from PIL import Image
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import docx  # python-docx for .docx CVs
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

load_dotenv(r"c:\Agent Coco\.env")
URL = os.environ["DATABASE_URL"]
HOST = URL.split("@")[1].split("/")[0]
SQL_ENDPOINT = f"https://{HOST}/sql"

OUTPUT_DIR = r"c:\Agent Coco\output\cv_texts_job41_new_batch"


def q(sql, params=None):
    r = requests.post(
        SQL_ENDPOINT,
        headers={"Neon-Connection-String": URL, "Content-Type": "application/json"},
        json={"query": sql, "params": params or []},
        timeout=120,
    )
    r.raise_for_status()
    return r.json()["rows"]


def parse_pdf(pdf_bytes):
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        text = ""
        for page in reader.pages:
            text += (page.extract_text() or "")
        if len(text.strip()) > 50:
            return text
    except Exception:
        pass
    if OCR_AVAILABLE:
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            ocr_text = ""
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text += pytesseract.image_to_string(img)
            doc.close()
            return ocr_text
        except Exception:
            pass
    return ""


def parse_docx(docx_bytes):
    if not DOCX_AVAILABLE:
        return ""
    try:
        d = docx.Document(io.BytesIO(docx_bytes))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def main():
    apps = q("""
        SELECT a.id AS app_id,
               TRIM(c.first_name || ' ' || COALESCE(c.last_name, '')) AS full_name,
               c.email, c.location, c.linkedin_url, c.resume_file_name,
               a.canned_answers->'desiredSalary'->>'answer'   AS desired_salary,
               a.canned_answers->'currentSalary'->>'answer'   AS current_salary,
               a.canned_answers->'lastSalary'->>'answer'      AS last_salary,
               a.canned_answers->'dateAvailable'->>'answer'   AS date_available,
               a.canned_answers->'address'->>'answer'         AS address,
               LENGTH(c.resume_data) AS resume_len
        FROM applications a
        JOIN candidates c ON c.id = a.candidate_id
        WHERE a.job_id = 41 AND a.status = 'new'
        ORDER BY a.id
    """)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    summary = []
    print(f"\n{'='*60}\nJob 41 GM-Karachi — new batch CV extraction ({len(apps)} apps)\n{'='*60}\n")

    for a in apps:
        app_id = a["app_id"]
        full_name = a["full_name"]
        fname = (a.get("resume_file_name") or "").lower()

        if not a.get("resume_len"):
            print(f"app {app_id:>5}  {full_name:<30} NO CV (LinkedIn stub)")
            summary.append({**a, "cv_len": 0, "readable": False, "note": "no_resume"})
            continue

        rows = q("SELECT c.resume_data FROM applications a JOIN candidates c ON c.id=a.candidate_id WHERE a.id = $1", [app_id])
        b64 = rows[0]["resume_data"] if rows else None
        try:
            raw = base64.b64decode(b64) if b64 else b""
        except Exception:
            raw = b""

        if fname.endswith(".docx"):
            cv_text = parse_docx(raw)
        else:
            cv_text = parse_pdf(raw)

        safe_name = re.sub(r'[^a-zA-Z0-9_]', '_', full_name)
        out_file = os.path.join(OUTPUT_DIR, f"{app_id}_{safe_name}.txt")
        with open(out_file, "w", encoding="utf-8") as f:
            f.write(f"APP ID:    {app_id}\n")
            f.write(f"NAME:      {full_name}\n")
            f.write(f"EMAIL:     {a['email']}\n")
            f.write(f"LOCATION:  {a.get('location') or 'Not mentioned'}\n")
            f.write(f"LINKEDIN:  {a.get('linkedin_url') or 'Not mentioned'}\n")
            f.write(f"DESIRED SALARY:  {a.get('desired_salary') or 'Not mentioned'}\n")
            f.write(f"CURRENT SALARY:  {a.get('current_salary') or a.get('last_salary') or 'Not mentioned'}\n")
            f.write(f"AVAILABLE: {a.get('date_available') or 'Not mentioned'}\n")
            f.write(f"ADDRESS:   {a.get('address') or 'Not mentioned'}\n")
            f.write("=" * 80 + "\n\n")
            f.write(cv_text if cv_text.strip() else "[CV UNREADABLE OR EMPTY]")

        ok = bool(cv_text.strip())
        print(f"app {app_id:>5}  {full_name:<30} {'OK  (' + str(len(cv_text)) + ' chars)' if ok else 'UNREADABLE'}")
        summary.append({k: a[k] for k in a if k != 'resume_data'} | {"cv_len": len(cv_text), "readable": ok, "file": out_file})

    with open(os.path.join(OUTPUT_DIR, "_summary.json"), "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)

    readable = sum(1 for s in summary if s["readable"])
    print(f"\n{'='*60}")
    print(f"DONE. {len(summary)} applications | readable {readable} | no-CV/unreadable {len(summary)-readable}")
    print(f"Output: {OUTPUT_DIR}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
