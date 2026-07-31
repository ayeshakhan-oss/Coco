# -*- coding: utf-8 -*-
"""Generate SMG + GM case study DOCX files mirroring the HOG case study layout,
then upload to Google Drive as native Google Docs."""
import re, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

LOGO = r"c:\Agent Coco\assets\logo_taleemabad.png"
OUTDIR = r"c:\Agent Coco\output\case_studies"
FONT = "Quicksand"
BLUE = RGBColor(0x3C, 0x78, 0xD8)
BLACK = RGBColor(0x00, 0x00, 0x00)

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

def add_runs(p, text, size=11, color=BLACK, italic=False, underline=False, bold_all=False):
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            style(r, size, color, False or bold_all, italic, underline)
        r = p.add_run(m.group(1))
        style(r, size, color, True, italic, underline)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        style(r, size, color, bold_all, italic, underline)

def style(run, size, color, bold, italic, underline):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    run.underline = underline
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(attr), FONT)

def build(doc_spec, path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        hdr = section.header
        hp = hdr.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.add_run().add_picture(LOGO, width=Inches(1.15))

    for item in doc_spec:
        kind = item[0]
        text = item[1] if len(item) > 1 else ""
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(6)
        if kind == "title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(18); pf.space_after = Pt(4)
            add_runs(p, text, size=17, bold_all=True)
        elif kind == "subtitle":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_after = Pt(16)
            add_runs(p, text, size=15, color=BLUE, bold_all=True)
        elif kind == "meta":
            pf.space_after = Pt(2)
            add_runs(p, text)
        elif kind == "h1":
            pf.space_before = Pt(16); pf.space_after = Pt(8)
            add_runs(p, text, size=15, bold_all=True)
        elif kind == "h2":
            pf.space_before = Pt(12); pf.space_after = Pt(6)
            add_runs(p, text, size=12.5, color=BLUE, bold_all=True)
        elif kind == "h3":
            pf.space_before = Pt(8); pf.space_after = Pt(4)
            add_runs(p, text, bold_all=True)
        elif kind == "h3u":
            pf.space_before = Pt(8); pf.space_after = Pt(4)
            add_runs(p, text, bold_all=True, underline=True)
        elif kind == "para":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, text)
        elif kind == "quote":
            pf.left_indent = Inches(0.4); pf.right_indent = Inches(0.2)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, text, italic=True)
        elif kind == "bullet":
            p.style = doc.styles["List Bullet"]
            pf.space_after = Pt(3)
            add_runs(p, text)
        elif kind == "bullet2":
            p.style = doc.styles["List Bullet 2"]
            pf.space_after = Pt(3)
            add_runs(p, text)
        elif kind == "num":
            p.style = doc.styles["List Number"]
            pf.space_after = Pt(3)
            add_runs(p, text)
        elif kind == "numpara":  # manually numbered paragraph with hanging text
            pf.left_indent = Inches(0.3)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, text)
        elif kind == "footer":
            pf.space_before = Pt(14); pf.space_after = Pt(0)
            add_runs(p, text, italic=True)
    doc.save(path)
    print("built", path)

SMG = [
    ("title", "Senior Manager Growth (SMG) @ Taleemabad"),
    ("subtitle", "Case Study: The Execution Sprint"),
    ("meta", "**Role:** Senior Manager Growth (SMG)"),
    ("meta", "**Time Recommended:** 2.5\u20133 hours total"),
    ("meta", "**Submission Format:** Written memo + slides + spreadsheet + short reflective response + Docs/PDFs (no video required)"),
    ("h1", "Role Overview"),
    ("para", "The **Senior Manager Growth** is a hands-on execution leader and the **second-in-command (2IC)** to the Head of Growth. This role translates growth strategy into action across acquisition, stakeholder engagement, product adoption, and pricing experiments. You will own defined growth initiatives, manage channels, execute growth loops, and close deals \u2014 while learning to think systemically using the Four Fits Model."),
    ("para", "This case study is designed to see how you **execute**: how you read data quickly, turn a strategy someone else designed into weekly action, and keep a deal moving when it stalls. We are not testing polish. We are testing judgment, follow-through, and clarity."),
    ("h1", "Ground Rules"),
    ("bullet", "Time-box yourself honestly. We calibrated this for 2.5\u20133 hours; a complete, rough answer beats a beautiful, partial one."),
    ("bullet", "You may use AI tools for any part of this. If you do, add one line at the end of each deliverable telling us how you used them. (The role expects you to work this way \u2014 we want to see how you do it.)"),
    ("bullet", "Do not fabricate data. If you make an assumption, label it as one."),
    ("h1", "Case Study Assignments for Candidates"),
    ("h2", "Assignment 1: Channel & Cohort Execution Analysis"),
    ("para", "**Context:** Alpha Platform (Dummy Name) is an AI-powered teaching assistant deployed via WhatsApp that helps educators improve their teaching practice. The platform offers:"),
    ("bullet", "**Coaching Sessions:** Teachers record classroom lessons and receive AI feedback"),
    ("bullet", "**Reading Assessments:** Students read passages aloud for fluency/accuracy analysis"),
    ("bullet", "**Lesson Plan Generation:** AI-assisted lesson planning"),
    ("bullet", "**Presentation Generation:** AI-assisted teaching presentations"),
    ("bullet", "**General Chat:** Free-form Q&A about teaching and pedagogy"),
    ("para", "**Your Data:** Please visit the Tab titled as **Dataset - Alpha Platform** to access the 6 weeks of anonymized user data (Nov 3 - Dec 16, 2025) from Alpha Platform's WhatsApp deployment across geographies. For this assignment, focus on:"),
    ("num", "01_master_user_dataset.csv - Comprehensive per-user metrics"),
    ("num", "06_daily_activity.csv - Aggregated daily platform metrics"),
    ("num", "07_country_breakdown.csv - User distribution by country"),
    ("num", "08_aggregate_metrics.csv - High-level platform summary"),
    ("para", "(You may pull from the other files in the dataset if useful, but the four above are sufficient.)"),
    ("h3u", "Your Analysis Must Address:"),
    ("h3", "Part 1: Where would you double down?"),
    ("bullet", "Which countries, user segments, and features show the strongest engagement and retention?"),
    ("bullet", "Which look like noise \u2014 activity that will not compound?"),
    ("bullet", "Pick the **two highest-leverage areas** for the next 8 weeks and defend the choice with numbers."),
    ("h3", "Part 2: Channel Experiments"),
    ("bullet", "Propose **three concrete experiments** to improve Product\u2013Channel Fit and Channel\u2013Model Fit for your chosen areas."),
    ("bullet", "For each experiment: hypothesis, what you would do in week 1, the metric that decides success, and the kill criterion (what result makes you stop)."),
    ("para", "**Deliverable:** Max 6 slides **or** a 3-page memo, plus one supporting spreadsheet showing your working. No appendix needed."),
    ("para", "**Data Access:** Here"),
    ("para", "**Dataset Description:** Here"),
    ("h2", "Assignment 2: Growth Loop Execution Plan"),
    ("para", "Growth leadership has designed the following loop for Alpha Platform, and it is now yours to run:"),
    ("quote", "A teacher uses Alpha Platform's coaching feature \u2192 receives feedback that visibly improves a lesson \u2192 the school administrator notices and asks about it \u2192 the administrator invites Taleemabad to run a session for the whole school \u2192 more teachers onboard \u2192 some of those teachers transfer schools or share in district WhatsApp groups \u2192 new schools ask in."),
    ("para", "Your job is **not** to redesign the loop. Your job is to make it spin."),
    ("h3", "Build a 60-day execution plan that covers:"),
    ("bullet", "The specific actions you would run at each step of the loop (who does what, through which channel \u2014 school visits, WhatsApp, events, direct outreach)."),
    ("bullet", "How you engage the different stakeholder layers (teachers, administrators, district officials) so the loop does not stall at any single layer."),
    ("bullet", "The metrics you track weekly, including how you would measure this loop's contribution toward a **K-factor of 0.2**."),
    ("bullet", "How you keep the pipeline visible: what you log, where, and what your weekly update to the Head of Growth contains."),
    ("bullet", "The two most likely points where the loop breaks, and your contingency for each."),
    ("para", "**Submit:** Max 2 pages + one simple tracker (spreadsheet) you would actually use to run this."),
    ("h2", "Assignment 3: The Stalled Deal"),
    ("para", "**Scenario:** Three months ago, Taleemabad ran a successful pilot with a district education authority: 40 schools, strong teacher usage, and the District Education Officer (DEO) publicly praised the results. The DEO verbally committed to expanding to 200 schools."),
    ("para", "Since then: six weeks of silence. Procurement says the file is \u201cin process.\u201d The DEO has stopped responding to calls. The provincial budget cycle closes in five weeks \u2014 if the expansion is not in this budget, it waits a year. Meanwhile, your Head of Growth is asking for a pipeline update, and a competitor has started meeting schools in the same district."),
    ("h3", "Your task \u2014 produce three things:"),
    ("num", "**Action plan (1 page):** Your moves for the next five weeks, in sequence. Who you contact, in what order, and what you ask each person for. Be specific about how you re-open a relationship that has gone quiet without damaging it."),
    ("num", "**The stakeholder email:** The actual email you would send to the DEO this week. (Write it in full \u2014 we want to see your judgment on tone, length, and ask.)"),
    ("num", "**Internal update (half page):** What you tell the Head of Growth \u2014 including your honest read on the probability this closes, and what help you need."),
    ("para", "**Submit:** One document containing all three parts."),
    ("h1", "One Last Ask (Before You Hit Send)"),
    ("para", "Add a short reflective response (**200 words or fewer**) sharing:"),
    ("bullet", "A time you spotted that a strategy was not working on the ground and flagged it early. What did you see, what did you do, and what happened?"),
    ("para", "Prefer speaking? A voice note (max 2 minutes) is equally welcome. Clarity > polish."),
    ("footer", "Taleemabad is an equal-opportunity employer."),
    ("footer", "We hire for grit, learning velocity, and outcomes."),
]

GM = [
    ("title", "Growth Manager @ Taleemabad"),
    ("subtitle", "Case Study: The Story, the Room, and the Deal"),
    ("meta", "**Role:** Growth Manager (Lahore)"),
    ("meta", "**Time Recommended:** 2\u20132.5 hours total"),
    ("meta", "**Submission Format:** Written memo + slides + spreadsheet + short reflective response + Docs/PDFs (no video required)"),
    ("h1", "Role Overview"),
    ("para", "The **Growth Manager** sits at the intersection of Growth, Strategic Engagement, and Storytelling. The role positions Taleemabad as a credible, impact-driven partner by shaping compelling narratives, leading high-level convenings, and building strategic relationships with governments and institutional stakeholders \u2014 converting interest into long-term collaborations and closed deals."),
    ("para", "This case study mirrors the three muscles the role uses every week: **telling the story** so a senior policymaker leans in, **designing the room** where ecosystem dialogue happens, and **moving a partnership** from conversation to closure."),
    ("h1", "Ground Rules"),
    ("bullet", "Time-box yourself honestly. We calibrated this for 2\u20132.5 hours; a complete, rough answer beats a beautiful, partial one."),
    ("bullet", "You may use AI tools for any part of this. If you do, add one line at the end of each deliverable telling us how you used them."),
    ("bullet", "Ground everything in the facts provided below or publicly verifiable information. If you make an assumption, label it as one."),
    ("h3", "What you may use about Taleemabad (all public):"),
    ("bullet", "Taleemabad began with learning cartoons broadcast on national TV, then built a learning app that brought literacy content to 1.5 million smartphones."),
    ("bullet", "It opened and ran 140 low-fee private schools."),
    ("bullet", "It now works inside Pakistan's public schooling system: two districts, thousands of teachers, with measurable learning gains."),
    ("bullet", "It has recently separated Growth and Fundraising; the Growth team builds strategic institutional partnerships with government, development partners, and sector leaders."),
    ("h1", "Case Study Assignments for Candidates"),
    ("h2", "Assignment 1: Strategic Storytelling"),
    ("para", "**Scenario:** Taleemabad has been invited to brief the **Secretary of School Education** of a province where it does not yet work. You have been asked to prepare the materials. The Secretary is smart, sceptical, short on time, and has seen a hundred edtech pitches that promised transformation and delivered tablets in cupboards."),
    ("h3", "Your task \u2014 produce two things:"),
    ("num", "**A policy-facing one-pager** the Secretary can read in three minutes: the problem, Taleemabad's approach, the evidence, and what a partnership could look like. It must be evidence-backed, emotionally resonant without being sentimental, and written for a policymaker \u2014 not a donor, not a customer."),
    ("num", "**A 5-slide narrative outline** for the in-person briefing (outline means: slide title + 2\u20133 bullets of what goes on each slide + one line on why it is there). We are looking at the arc of your story, not your slide design."),
    ("para", "**Submit:** One document containing both."),
    ("h2", "Assignment 2: The Room \u2014 Designing a Strategic Convening"),
    ("para", "**Scenario:** Taleemabad has completed a study showing measurable learning gains from its work in public schools. Leadership wants to launch this study as a moment that positions Taleemabad as a thought leader in education transformation \u2014 **not** a marketing event."),
    ("h3", "Design the convening. Your event brief must cover:"),
    ("bullet", "**Objective:** What must be true the morning after for this event to have been worth it?"),
    ("bullet", "**The room:** Who is invited and why \u2014 map 8\u201312 attendee types (e.g., secretaries, senior government officials, development partners, sector leaders) and what each should walk away thinking."),
    ("bullet", "**The agenda:** Structure of the session(s), including how you keep senior stakeholders engaged rather than speech-fatigued."),
    ("bullet", "**Three talking points** for Taleemabad's leadership on stage."),
    ("bullet", "**The follow-up engine:** How engagement converts after the event \u2014 who gets what within 7 days, and how conversations become pipeline."),
    ("para", "**Submit:** Max 6 slides **or** 4 pages."),
    ("h2", "Assignment 3: The Pipeline \u2014 Market Scan & Deal Path"),
    ("para", "**Part 1: Market scan.** Identify **five real prospective institutional partners** for Taleemabad in Punjab or at the federal level \u2014 across government and the development sector. For each: who they are, why they fit, and what the opening conversation is. (Use real organizations and public information; label any assumptions.)"),
    ("para", "**Part 2: Pipeline structure.** Lay out the five in a simple pipeline tracker (spreadsheet): engagement stage, next action, owner, and what \u201cprogress\u201d means for each in the next 30 days."),
    ("para", "**Part 3: One deal, end to end.** Pick the single most promising partner and map the path from first conversation to closed deal: the likely stages, the decision-makers involved, where it will probably stall, and how you keep momentum through a long institutional cycle."),
    ("para", "**Submit:** Spreadsheet + max 2-page note."),
    ("h1", "One Last Ask (Before You Hit Send)"),
    ("para", "Add a short reflective response (**200 words or fewer**) sharing:"),
    ("bullet", "A room you led \u2014 a discussion, briefing, or convening where you moderated senior stakeholders. What made it work, and what would you do differently?"),
    ("para", "Prefer speaking? A voice note (max 2 minutes) is equally welcome. Clarity > polish."),
    ("footer", "Taleemabad is an equal-opportunity employer."),
    ("footer", "We hire for grit, learning velocity, and outcomes."),
]

os.makedirs(OUTDIR, exist_ok=True)
smg_path = os.path.join(OUTDIR, "SMG_case_study.docx")
gm_path = os.path.join(OUTDIR, "GM_case_study.docx")
build(SMG, smg_path)
build(GM, gm_path)

# ---- Upload to Google Drive as native Google Docs ----
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build as gbuild
from googleapiclient.http import MediaFileUpload

creds = Credentials.from_authorized_user_file(r"c:\Agent Coco\.claude\config\token_sheets_broad.json")
drive = gbuild("drive", "v3", credentials=creds)

for path, title in [
    (smg_path, "Senior Manager Growth (SMG) @ Taleemabad \u2014 Case Study: The Execution Sprint"),
    (gm_path, "Growth Manager @ Taleemabad \u2014 Case Study: The Story, the Room, and the Deal"),
]:
    media = MediaFileUpload(path, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    f = drive.files().create(
        body={"name": title, "mimeType": "application/vnd.google-apps.document"},
        media_body=media,
        fields="id, webViewLink",
    ).execute()
    print(title)
    print("  ->", f["webViewLink"])
