#!/usr/bin/env python3
"""
Contract / NDA .docx validator — Skill 07 eval harness.

Every check here encodes a defect that actually reached Ayesha in a pilot.
Source session: Muhammad Shayan Fellow package, 2026-08-13 (three review rounds).

Usage:
    python scripts/evals/contract_docx_eval.py --file "<path.docx>" --type fellow
    python scripts/evals/contract_docx_eval.py --dir "output/contracts/Muhammad Shayan" --type fellow

Exit codes:
    0 = clean, or warnings only
    2 = HARD BLOCK violations present
"""

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx not installed — cannot validate", file=sys.stderr)
    sys.exit(0)

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# ── Rules learned the hard way ────────────────────────────────────────────────

# Any of these left in a finished document means a field was never filled.
PLACEHOLDERS = [
    "XYZ", "X Y Z", "EMPLOYEE NAME", "EMPLOYEE'S NAME", "EMPLOYEE’S NAME",
    "EMPLOYEE'S CNIC", "EMPLOYEE’S CNIC", "FELLOW NAME", "JOINING DATE",
    "CURRENT DATE", "Current Date", "DATE, MONTH", "EMPLOYER NAME",
    "DESIGNATION", "(joining date)", "Mention Job Description here",
]

# "Mr./Mrs." is a drafting option, not contract wording — Ayesha 2026-08-13.
SALUTATION_ARTEFACT = "Mr./Mrs."

# Clauses Ayesha removed for Fellow engagements (2026-08-13).
FELLOW_FORBIDDEN = [
    ("travel incurred for business purposes", "business-travel clause"),
    ("technical equipment", "technical equipment / asset-damage clause"),
    ("In-patient (IPD) Medical Coverage", "IPD medical coverage clause"),
    ("No travel or daily allowances", "no-travel-or-daily-allowances clause"),
    ("probationary 13 weeks", "13-week / 3-month probation clause"),
    ("Unlimited trust-based leaves", "unlimited trust-based leave clause"),
    ("Base Salary: PKR", "Base Salary line (Fellows get Total Earnings only)"),
    ("Medical: PKR", "Medical line (Fellows get Total Earnings only)"),
    ("Others: PKR", "Others line (Fellows get Total Earnings only)"),
]

# The replacement leave terms that MUST be present on a Fellow contract.
FELLOW_REQUIRED = [
    ("annual leave of 3 working days", "Fellow annual-leave entitlement"),
    ("medical leave of 3 working days", "Fellow medical-leave entitlement"),
    ("maximum limit of 03 days", "medical-leave sub-condition"),
    ("Only unpaid leaves can be granted", "unpaid-leave sub-condition"),
]


def _paragraph_texts(doc):
    out = [(f"para {i}", p.text) for i, p in enumerate(doc.paragraphs)]
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                out.append((f"table {ti} r{ri}c{ci}", cell.text))
    return out


def evaluate_contract(path: Path, doc_type: str):
    """Return (blocks, warnings) for a populated contract or NDA."""
    doc = Document(path)
    blocks, warns = [], []
    is_fellow = doc_type == "fellow"
    is_nda = "nda" in path.name.lower()

    texts = _paragraph_texts(doc)
    joined = "\n".join(t for _, t in texts)

    # 1 — unresolved placeholders
    for ph in PLACEHOLDERS:
        for loc, txt in texts:
            if ph in txt:
                blocks.append(f"unresolved placeholder {ph!r} at {loc}")
                break

    # 2 — leftover yellow highlighting
    highlights = 0
    for p in doc.paragraphs:
        highlights += sum(1 for r in p.runs if r.font.highlight_color is not None)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    highlights += sum(
                        1 for r in p.runs if r.font.highlight_color is not None
                    )
    if highlights:
        blocks.append(f"{highlights} run(s) still carry yellow highlighting")

    # 3 — unresolved salutation option
    if SALUTATION_ARTEFACT in joined:
        blocks.append(
            f"{SALUTATION_ARTEFACT!r} left in the document — pick the actual salutation"
        )

    # 4 — forced page breaks between clauses (the "big gaps" defect)
    #     In-body sectPr set to NEW_PAGE starts a fresh page mid-contract.
    new_page_sections = 0
    for i, section in enumerate(doc.sections[:-1]):
        if int(section.start_type) == 2:  # WD_SECTION.NEW_PAGE
            new_page_sections += 1
    if new_page_sections:
        blocks.append(
            f"{new_page_sections} in-body section break(s) still set to NEW_PAGE — "
            f"these force blank gaps between points; set them to CONTINUOUS"
        )

    # 5 — stray empty auto-numbered paragraph (renders as a lone orphan bullet)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            continue
        pPr = p._p.find(W + "pPr")
        if pPr is not None and pPr.find(W + "numPr") is not None:
            blocks.append(
                f"empty auto-numbered paragraph at para {i} — renders as a stray "
                f"bullet; strip its numbering (do not delete, it may hold sectPr)"
            )

    # 6 — Annexure-A narrow-column defect
    for i, p in enumerate(doc.paragraphs):
        stripped = p.text.strip()
        if stripped.startswith("Job Description:") or stripped.startswith("•"):
            ri_ = p.paragraph_format.right_indent
            if ri_ and int(ri_) > 457200:  # > 0.5"
                blocks.append(
                    f"para {i} has a {int(ri_)/914400:.1f}\" right indent — this "
                    f"collapses Annexure A into a narrow column"
                )

    # 7 — Annexure heading must break onto two lines
    if not is_nda:
        for p in doc.paragraphs:
            if p.text.strip().startswith("Job Description:") and "Key Responsibilities" in p.text:
                if not p._p.findall(f".//{W}br"):
                    blocks.append(
                        "Annexure-A heading must break onto two lines: "
                        "'Job Description:' / 'Key Responsibilities:'"
                    )
                break

    # 7b — headings must be explicitly bold and glued to their content.
    #      Style-inherited bold does not survive Drive's PDF conversion, and a
    #      heading without keep_with_next strands at the foot of a page.
    if not is_nda:
        unbolded, unglued = [], []
        paras = doc.paragraphs
        for i, p in enumerate(paras):
            txt = p.text.strip()
            runs = [r for r in p.runs if r.text.strip()]
            if not txt or txt == "AND" or not runs:
                continue
            is_heading = p.style.name.startswith("Heading")
            if not is_heading:
                continue
            if not all(r.bold for r in runs):
                unbolded.append(f"para {i} ({txt[:32]})")
            has_following = any(q.text.strip() for q in paras[i + 1:])
            if has_following and not p.paragraph_format.keep_with_next:
                unglued.append(f"para {i} ({txt[:32]})")
        if unbolded:
            blocks.append(
                "heading(s) not explicitly bold — style-inherited bold is lost in "
                f"PDF conversion: {', '.join(unbolded[:4])}"
            )
        if unglued:
            blocks.append(
                "heading(s) missing keep_with_next — they can strand at a page "
                f"foot: {', '.join(unglued[:4])}"
            )

    # 7c — header labels Date: / CNIC: / Name: must be bold
    if not is_nda:
        for p in doc.paragraphs[:15]:
            if not p.runs:
                continue
            key = p.runs[0].text.strip().rstrip(":")
            if key in ("Date", "CNIC", "Name") and not p.runs[0].bold:
                blocks.append(f"header label {key!r} is not bold")

    # 8 — Fellow-specific clause rules
    if is_fellow and not is_nda:
        for needle, label in FELLOW_FORBIDDEN:
            if needle.lower() in joined.lower():
                blocks.append(f"{label} must be REMOVED for Fellows (found {needle!r})")
        for needle, label in FELLOW_REQUIRED:
            if needle.lower() not in joined.lower():
                blocks.append(f"missing {label} (expected text: {needle!r})")

        # 9 — inherited NIETE project line
        if "National Institute of Excellence in Teacher Education" in joined:
            warns.append(
                "Project line still reads 'National Institute of Excellence in "
                "Teacher Education' — inherited from the NIETE master. Confirm with "
                "Ayesha unless this really is a NIETE engagement."
            )

    # 10 — NDA must carry nothing but name + dates
    if is_nda:
        import re

        if re.search(r"\d{5}-\d{7}-\d", joined):
            blocks.append("NDA contains a CNIC — NDAs take name + dates ONLY")

    return blocks, warns


# ── Joining-email rules (locked verbatim by Ayesha 2026-08-13) ────────────────

# The onboarding form differs by programme — sending the wrong one is a real error.
FELLOW_FORM = "docs.google.com/forms/d/e/1FAIpQLSf70SM4jlx4muDMLlN1ZMqHqVEQjJQgCBga-oRM-M1OZXCePw"
NIETE_FORM = "docs.google.com/forms/d/e/1FAIpQLSdVAYfCZZhusF_tNLn7mxzoK5BFXDa7xfj2FZifRlva-YDBHQ"
WHATSAPP_GROUP = "chat.whatsapp.com/HglkfuENmLqEbaq8N5jSVq"


def evaluate_joining_email(body: str, attachments: list):
    """Check a Fellow joining email against the locked templates.

    `attachments` is a list of filenames actually being sent.
    """
    import re

    blocks, warns = [], []
    # Collapse Python/HTML string-concatenation artefacts ("abc"\n  "def") so a
    # URL split across source lines is still recognised as one literal.
    body = re.sub(r"['\"]\s*\n\s*['\"]", "", body)
    low = body.lower()

    has_contract = any("contract - " in a.lower() for a in attachments)
    has_nda = any(a.lower().startswith("nda - ") or "nda - " in a.lower()
                  for a in attachments)

    # Candidate-facing documents go out as PDF, never Word (Ayesha 2026-08-13).
    for a in attachments:
        if a.lower().endswith(".docx"):
            blocks.append(
                f"attachment {a!r} is a Word document — contracts and NDAs must be "
                f"sent as PDF (convert with scripts/utils/docx_to_pdf_drive.py)"
            )

    is_volunteer = "volunteer fellowship" in low
    is_transition = "transition to a paid fellowship" in low

    # THE package rule — a volunteer Fellow never receives a contract.
    if is_volunteer and has_contract:
        blocks.append(
            "VOLUNTEER fellowship email is attaching a CONTRACT — volunteer/unpaid "
            "Fellows get the NDA ONLY (Ayesha 2026-08-13)"
        )
    if is_volunteer and not has_nda:
        blocks.append("volunteer fellowship email must attach the NDA")

    # Transition email: contract only, NDA already signed at the unpaid stage.
    if is_transition and has_nda:
        blocks.append(
            "unpaid-to-paid transition email is attaching an NDA — the NDA was "
            "already signed at the unpaid stage; send the Contract only"
        )

    # Paid new-joiner: both documents.
    if not is_volunteer and not is_transition and "fellowship" in low:
        if has_contract and not has_nda:
            warns.append("paid fellowship email attaches a contract but no NDA")

    # No pilot/meta commentary inside a candidate-facing email body. The pilot and
    # the live email must be byte-identical except for the recipient — notes to
    # Ayesha belong in chat, not in the email (Ayesha 2026-08-13).
    # Narrow, unambiguous phrases only — these can never appear in a real
    # candidate email, so they will not false-positive on code comments.
    META_MARKERS = [
        "nothing has gone to", "on live send", "still open:",
        "two things to check", "pilot preface", "for your review before",
    ]
    for marker in META_MARKERS:
        if marker in low:
            blocks.append(
                f"email body contains pilot/meta commentary ({marker!r}) — the pilot "
                f"must be identical to what the candidate receives; put notes in chat"
            )
            break

    # Never name the weekday — the date alone (Ayesha 2026-08-13).
    WEEKDAYS = ("monday", "tuesday", "wednesday", "thursday", "friday",
                "saturday", "sunday")
    for day in WEEKDAYS:
        if re.search(rf"\b{day}\b,?\s*\d", low):
            blocks.append(
                f"email names the weekday ({day.title()}) before a date — write the "
                f"date only, no day name"
            )
            break

    # Key details must be bold: joining/effective date, compensation, duration.
    # These are what the candidate scans for (Ayesha 2026-08-13).
    # Emphasis counts whether it comes from <b>/<strong> or a font-weight:bold
    # style — both render bold, and table-based designs use the latter.
    bolded = " ".join(
        re.findall(r"<(?:b|strong)[^>]*>(.*?)</(?:b|strong)>", body, flags=re.S | re.I)
        + re.findall(r"<[^>]*font-weight:\s*bold[^>]*>(.*?)<", body, flags=re.S | re.I)
    ).lower()
    # Trigger on the presence of the details themselves, not on particular words —
    # a narrower trigger silently skipped the check on emails that never said
    # "fellowship" or "compensation".
    if re.search(r"pkr\s*[\d,]+", low) or "click here" in low:
        money = re.search(r"pkr\s*[\d,]+", low)
        if money and money.group(0) not in bolded:
            blocks.append(
                f"compensation ({money.group(0).upper()}) is not bold — always bold "
                f"the joining date, compensation and duration"
            )
        date_m = re.search(
            r"\d{1,2}(?:st|nd|rd|th)?\s+(?:of\s+)?"
            r"(?:january|february|march|april|may|june|july|august|september|"
            r"october|november|december)\s+\d{4}", low)
        if date_m and date_m.group(0) not in bolded:
            blocks.append(
                f"joining/effective date ({date_m.group(0)}) is not bold — always "
                f"bold the joining date, compensation and duration"
            )

    # Locked links must be present and hyperlinked, never bare.
    if "click here" in low:
        is_niete = "niete" in low
        has_fellow_form = FELLOW_FORM in body
        has_niete_form = NIETE_FORM in body

        if not (has_fellow_form or has_niete_form):
            blocks.append("joining email is missing the locked onboarding form link")
        elif is_niete and not has_niete_form:
            blocks.append(
                "NIETE joining email is using the FELLOW onboarding form — NIETE "
                "hires get the NIETE form"
            )
        elif not is_niete and not has_fellow_form:
            blocks.append(
                "non-NIETE joining email is using the NIETE onboarding form"
            )
        if WHATSAPP_GROUP not in body:
            blocks.append("joining email is missing the locked WhatsApp group link")

    return blocks, warns


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--file")
    ap.add_argument("--dir")
    ap.add_argument("--type", default="fellow",
                    choices=["fellow", "fulltime", "project", "addendum"])
    args = ap.parse_args()

    targets = []
    if args.file:
        targets.append(Path(args.file))
    if args.dir:
        targets.extend(sorted(Path(args.dir).glob("*.docx")))
    if not targets:
        ap.error("pass --file or --dir")

    total_blocks = 0
    for t in targets:
        if not t.exists():
            print(f"MISSING: {t}")
            total_blocks += 1
            continue
        blocks, warns = evaluate_contract(t, args.type)
        total_blocks += len(blocks)
        status = "BLOCKED" if blocks else ("WARN" if warns else "PASS")
        print(f"\n[{status}] {t.name}")
        for b in blocks:
            print(f"   HARD BLOCK: {b}")
        for w in warns:
            print(f"   warning:    {w}")
        if not blocks and not warns:
            print("   all checks passed")

    print(f"\n{'=' * 60}")
    if total_blocks:
        print(f"RESULT: {total_blocks} hard block(s) — do NOT send")
        return 2
    print("RESULT: clean")
    return 0


if __name__ == "__main__":
    sys.exit(main())
