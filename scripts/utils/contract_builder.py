"""
Shared contract-document builder — Skill 07.

Every helper here exists because a defect reached Ayesha in a pilot during the
Muhammad Shayan Fellow package (2026-08-13, six rounds). Any new contract, NDA or
addendum builder MUST use these rather than hand-rolling docx edits, so the same
mistakes cannot recur.

Typical use:

    from scripts.contracts.contract_builder import *

    doc, p = open_master(MASTER, OUT)
    drop = [p[30], p[31]]                      # capture objects BEFORE editing
    fill_paragraph(p[8], ["Full Name"])
    bold_header_labels(p)                      # Date: / CNIC: / Name:
    apply_heading_formatting(doc)
    make_sections_continuous(doc)
    for q in drop: delete_paragraph(q)
    doc.save(OUT)
    pdf = to_pdf(OUT)                          # candidates receive PDF, never .docx
"""

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.text.paragraph import Paragraph

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

__all__ = [
    "open_master", "highlighted_segments", "fill_segment", "fill_paragraph",
    "strip_numbering", "strip_section_break", "delete_paragraph",
    "add_bullet_after", "clone_paragraph_after", "add_hyperlink", "delete_table_row",
    "apply_heading_formatting", "bold_header_labels",
    "make_sections_continuous", "clear_trailing_orphan_bullet", "to_pdf",
]


# ── Setup ─────────────────────────────────────────────────────────────────────

def open_master(master: Path, out: Path):
    """Copy a master to the output path and open the COPY. Masters are never edited."""
    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(master, out)
    doc = Document(out)
    return doc, doc.paragraphs


# ── Fill fields ───────────────────────────────────────────────────────────────

def highlighted_segments(paragraph):
    """Group contiguous highlighted runs into segments: [[run, run], [run], ...]."""
    segments, current = [], []
    for run in paragraph.runs:
        if run.font.highlight_color is not None:
            current.append(run)
        elif current:
            segments.append(current)
            current = []
    if current:
        segments.append(current)
    return segments


def fill_segment(runs, value):
    """Write `value` into a segment; run 0 takes the text, the rest are emptied.

    Preserves the placeholder's leading/trailing whitespace so sentence spacing is
    unchanged, and clears the yellow highlight.
    """
    original = "".join(r.text for r in runs)
    lead = original[: len(original) - len(original.lstrip())]
    trail = original[len(original.rstrip()):]
    runs[0].text = f"{lead}{value}{trail}"
    for r in runs[1:]:
        r.text = ""
    for r in runs:
        r.font.highlight_color = None


def fill_paragraph(paragraph, values):
    """Fill each highlighted segment of a paragraph, in document order.

    Raises if the count does not match — a silent mismatch means a field was
    filled into the wrong place.
    """
    segments = highlighted_segments(paragraph)
    if len(segments) != len(values):
        raise AssertionError(
            f"expected {len(values)} highlighted segment(s), found {len(segments)} "
            f"in: {paragraph.text[:80]!r}"
        )
    for seg, val in zip(segments, values):
        fill_segment(seg, val)


# ── XML hygiene ───────────────────────────────────────────────────────────────

def strip_numbering(paragraph):
    """Remove auto-list numbering (kills the bullet glyph, keeps the paragraph)."""
    pPr = paragraph._p.find(W_NS + "pPr")
    if pPr is not None:
        numPr = pPr.find(W_NS + "numPr")
        if numPr is not None:
            pPr.remove(numPr)


def strip_section_break(paragraph):
    """Remove a section break carried in a paragraph's properties.

    Critical when cloning: several master paragraphs carry a NEW_PAGE sectPr and a
    deepcopy duplicates it, so every clone starts its own page.
    """
    pPr = paragraph._p.find(W_NS + "pPr")
    if pPr is not None:
        sectPr = pPr.find(W_NS + "sectPr")
        if sectPr is not None:
            pPr.remove(sectPr)


def delete_paragraph(paragraph):
    """Remove a paragraph entirely. Capture targets as objects BEFORE any edits."""
    el = paragraph._element
    el.getparent().remove(el)


def add_hyperlink(paragraph, text, url, template_run=None, replace=True):
    """Add a real Word hyperlink (blue, underlined) to a paragraph.

    python-docx has no native hyperlink API — this creates the w:hyperlink element
    and the external relationship by hand.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if replace:
        for r in list(paragraph.runs):
            r._element.getparent().remove(r._element)

    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if template_run is not None and template_run.font.name:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), template_run.font.name)
        rFonts.set(qn("w:hAnsi"), template_run.font.name)
        rPr.append(rFonts)
    if template_run is not None and template_run.font.size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(template_run.font.size.pt * 2)))
        rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return paragraph


def clone_paragraph_after(anchor, text, template_run=None, keep_numbering=True):
    """Insert a paragraph after `anchor`, inheriting its look.

    Unlike `add_bullet_after`, this KEEPS the anchor's auto-numbering so a new
    clause joins the existing numbered list. Section breaks are always stripped.
    """
    new_p = copy.deepcopy(anchor._p)
    for r in new_p.findall(W_NS + "r"):
        new_p.remove(r)
    for h in new_p.findall(W_NS + "hyperlink"):
        new_p.remove(h)
    anchor._p.addnext(new_p)

    para = Paragraph(new_p, anchor._parent)
    strip_section_break(para)
    if not keep_numbering:
        strip_numbering(para)

    src = template_run if template_run is not None else (
        anchor.runs[0] if anchor.runs else None
    )
    run = para.add_run(text)
    if src is not None:
        run.font.name = src.font.name
        run.font.size = src.font.size
    run.font.bold = False
    return para


def delete_table_row(table, row):
    """Remove an entire row (label + value) from a table."""
    row._element.getparent().remove(row._element)


def add_bullet_after(anchor, text, template_run, left_indent=228600, marker="•  "):
    """Insert a manually-bulleted paragraph directly after `anchor`.

    The clone inherits the anchor's look, then has its runs, auto-numbering and any
    section break stripped — clones carry baggage.
    """
    new_p = copy.deepcopy(anchor._p)
    for r in new_p.findall(W_NS + "r"):
        new_p.remove(r)
    anchor._p.addnext(new_p)

    para = Paragraph(new_p, anchor._parent)
    strip_numbering(para)
    strip_section_break(para)

    run = para.add_run(f"{marker}{text}")
    run.font.name = template_run.font.name
    run.font.size = template_run.font.size
    run.font.bold = False
    run.font.underline = False

    pf = para.paragraph_format
    pf.left_indent = left_indent
    pf.right_indent = 0        # masters carry a ~5.3" right indent that clones inherit
    pf.first_line_indent = 0
    pf.space_after = 60000
    return para


# ── Layout ────────────────────────────────────────────────────────────────────

def make_sections_continuous(doc):
    """Convert in-body NEW_PAGE section breaks to CONTINUOUS.

    Masters carry several NEW_PAGE section breaks that strand single lines on
    otherwise blank pages. All sections share the same page setup, so this removes
    the gaps without changing page geometry.
    """
    changed = 0
    for section in doc.sections[:-1]:
        if int(section.start_type) != int(WD_SECTION.CONTINUOUS):
            section.start_type = WD_SECTION.CONTINUOUS
            changed += 1
    return changed


def clear_trailing_orphan_bullet(doc):
    """Strip numbering from empty auto-numbered paragraphs (stray lone bullets).

    Never deletes them — they may carry section properties.
    """
    fixed = 0
    for p in doc.paragraphs:
        if p.text.strip():
            continue
        pPr = p._p.find(W_NS + "pPr")
        if pPr is not None and pPr.find(W_NS + "numPr") is not None:
            strip_numbering(p)
            p.paragraph_format.left_indent = 0
            p.paragraph_format.first_line_indent = 0
            fixed += 1
    return fixed


def bold_header_labels(paras, labels=("Date", "CNIC", "Name"), search_limit=15):
    """Bold the header labels (Date: / CNIC: / Name:), leaving values normal weight."""
    done = []
    for p in paras[:search_limit]:
        if not p.runs:
            continue
        first = p.runs[0]
        key = first.text.strip().rstrip(":")
        if key in labels:
            first.bold = True
            done.append(key)
    missing = set(labels) - set(done)
    if missing:
        raise AssertionError(f"header labels not found: {sorted(missing)}")
    return done


def apply_heading_formatting(doc, max_block=10, skip_texts=("AND",), bold=True):
    """Bold every section heading and keep it with the block that follows.

    Fixes two defects:
      1. Heading runs inherit bold from the style; Drive's PDF conversion flattens
         that, so bold must be set on the run.
      2. Without keep_with_next a heading strands at the foot of a page while its
         content starts the next one.

    `max_block` caps how much is glued to a heading — gluing a whole long section
    shunts it to a new page and re-creates the blank-gap problem.
    """
    paras = doc.paragraphs

    def looks_like_heading(p):
        """Short, fully-bold, non-sentence lines styled 'normal' are headings in
        everything but name (e.g. 'Contract Period & Compensation')."""
        txt = p.text.strip()
        runs = [r for r in p.runs if r.text.strip()]
        return (
            0 < len(txt) < 60
            and not txt.endswith(".")
            and bool(runs)
            and all(r.bold for r in runs)
        )

    count = 0
    for i, p in enumerate(paras):
        txt = p.text.strip()
        if not txt or txt in skip_texts:
            continue
        if not (p.style.name.startswith("Heading") or looks_like_heading(p)):
            continue

        # `bold=False` is for NDAs: fix the page flow without altering the
        # approved document's appearance.
        if bold:
            for run in p.runs:
                run.bold = True
        p.paragraph_format.keep_with_next = True
        count += 1

        # Build the CHAIN from the heading to the end of its block, including any
        # empty spacer paragraphs. keep_with_next only binds a paragraph to the
        # very next one, so a blank spacer between heading and content breaks the
        # chain and the content still falls to the next page — this is exactly how
        # "OFFER ACCEPTANCE:" ended up alone at the foot of page 5.
        # The block runs to the NEXT heading (or the cap) — not to the first blank
        # line. A section's own sub-points are often separated by spacers, and
        # "OFFER ACCEPTANCE:" must stay with its signature lines, not just its
        # first sentence.
        chain, content_seen = [], 0
        for q in paras[i + 1:]:
            txt = q.text.strip()
            if txt and (q.style.name.startswith("Heading") or looks_like_heading(q)):
                break              # next section starts here
            chain.append(q)
            if txt:
                content_seen += 1
                if content_seen >= max_block:
                    break
        # drop trailing blanks so we do not glue the next section on
        while chain and not chain[-1].text.strip():
            chain.pop()
        # every link except the last must point at the next one
        for q in chain[:-1]:
            q.paragraph_format.keep_with_next = True
        if chain:
            chain[-1].paragraph_format.keep_with_next = False

    return count


# ── Output ────────────────────────────────────────────────────────────────────

def to_pdf(docx_path):
    """Convert to PDF. Candidates receive PDF, never .docx.

    Delegates to the Drive converter — the only renderer available on this machine.
    Drive re-flows the document, so the result must still be eyeballed by Ayesha.
    """
    import sys

    sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "utils"))
    from docx_to_pdf_drive import convert

    return convert(Path(docx_path))
